import os
from datetime import datetime

import streamlit as st
import vertexai
from vertexai.generative_models import (
    GenerationConfig,
    GenerativeModel,
    HarmBlockThreshold,
    HarmCategory,
    Part,
)
from docx import Document
import re

PROJECT_ID = os.environ.get("GCP_PROJECT")  # Your Google Cloud Project ID
LOCATION = os.environ.get("GCP_REGION")  # Your Google Cloud Project Region
vertexai.init(project=PROJECT_ID, location=LOCATION)

min_date = datetime(1900, 1, 1)
max_date = datetime.now()

@st.cache_resource
def load_models():
    text_model_pro = GenerativeModel("gemini-1.0-pro")
    multimodal_model_pro = GenerativeModel("gemini-1.0-pro-vision")
    return text_model_pro, multimodal_model_pro


def get_gemini_pro_text_response(
    model: GenerativeModel,
    contents: str,
    generation_config: GenerationConfig,
    stream: bool = True,
):
    safety_settings = {
        HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
    }

    responses = model.generate_content(
        contents,
        generation_config=generation_config,
        safety_settings=safety_settings,
        stream=stream,
    )

    final_response = []
    for response in responses:
        try:
            final_response.append(response.text)
        except IndexError:
            final_response.append("")
            continue
    return " ".join(final_response)


# Set page configuration
st.set_page_config(
    page_title="Underwriting Risk Assessment Report Generator",
    page_icon=":house:",
)

# Custom CSS for styling
st.markdown("""
    <style>
    .main {
        background-color: #E9F1FA;
        padding: 20px;
    }
    .report-box {
        background-color: white;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 2px 2px 10px rgba(0, 0, 0, 0.1);
    }
    .header-title {
        color: #1D1D1;
        font-weight: bold;
        text-align: center;
    }
    .section-title {
        color: #333333;
        font-size: 18px;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    .stButton>button {
        background-color: #00ABE4;
        color: white;
    }
    </style>
""", unsafe_allow_html=True)

# Header
st.markdown("<h1 class='header-title'>Underwriting Risk Assessment Report Generator</h1>", unsafe_allow_html=True)

text_model_pro, multimodal_model_pro = load_models()

# Input Form
with st.form("input_form"):
    st.subheader("Applicant Information")
    applicant_name = st.text_input("Enter applicant name:", key="applicant_name", value="John Doe")
    date_of_birth = st.date_input("Enter date of birth:", key="date_of_birth", min_value=min_date, max_value=max_date)
    ssn = st.text_input("Enter Social Security Number:", key="ssn", value="123-45-6789")
    marital_status = st.selectbox("Select marital status:", ["Single", "Married", "Divorced", "Widowed"], key="marital_status")
    number_of_dependents = st.number_input("Enter number of dependents:", min_value=0, key="number_of_dependents", value=2)

    st.subheader("Employment and Income")
    current_employer = st.text_input("Enter current employer:", key="current_employer", value="XYZ Corporation")
    job_title = st.text_input("Enter job title:", key="job_title", value="Software Engineer")
    years_at_current_job = st.number_input("Enter years at current job:", min_value=0.0, key="years_at_current_job", value=5.0)
    monthly_gross_income = st.number_input("Enter monthly gross income:", min_value=0.0, key="monthly_gross_income", value=10000.0)
    other_income_sources = st.text_input("Enter other income sources:", key="other_income_sources", value="None")

    st.subheader("Financial Information")
    current_assets = st.number_input("Enter current assets:", min_value=0.0, key="current_assets", value=150000.0)
    current_debts = st.number_input("Enter current debts:", min_value=0.0, key="current_debts", value=50000.0)
    monthly_debt_obligations = st.number_input("Enter monthly debt obligations:", min_value=0.0, key="monthly_debt_obligations", value=1500.0)
    credit_score = st.number_input("Enter credit score:", min_value=0, max_value=850, key="credit_score", value=720)

    st.subheader("Property Information")
    property_address = st.text_input("Enter property address:", key="property_address", value="123 Elm Street, Springfield")
    property_type = st.selectbox("Select property type:", ["Single-family home", "Multi-family home", "Condo", "Townhouse"], key="property_type")
    purchase_price = st.number_input("Enter purchase price:", min_value=0.0, key="purchase_price", value=400000.0)
    loan_amount_requested = st.number_input("Enter loan amount requested:", min_value=0.0, key="loan_amount_requested", value=320000.0)
    down_payment_amount = st.number_input("Enter down payment amount:", min_value=0.0, key="down_payment_amount", value=80000.0)

    st.subheader("Loan Details")
    loan_type = st.selectbox("Select loan type:", ["Fixed-rate", "Adjustable-rate"], key="loan_type")
    loan_term = st.number_input("Enter loan term (years):", min_value=0, key="loan_term", value=30)
    interest_rate = st.number_input("Enter interest rate (%):", min_value=0.0, key="interest_rate", value=3.5)

    st.subheader("Housing Expense Information")
    current_monthly_housing_payment = st.number_input("Enter current monthly housing payment:", min_value=0.0, key="current_monthly_housing_payment", value=2000.0)
    estimated_property_taxes = st.number_input("Enter estimated property taxes (annual):", min_value=0.0, key="estimated_property_taxes", value=4000.0)
    homeowners_insurance = st.number_input("Enter homeowners insurance (annual):", min_value=0.0, key="homeowners_insurance", value=1200.0)
    hoa_fees = st.number_input("Enter HOA fees (monthly):", min_value=0.0, key="hoa_fees", value=0.0)

    st.subheader("Additional Information")
    purpose_of_loan = st.text_input("Enter purpose of loan:", key="purpose_of_loan", value="Purchase")
    bankruptcy_history = st.selectbox("Select bankruptcy history:", ["None", "Within the past year", "1-5 years ago", "More than 5 years ago"], key="bankruptcy_history")
    foreclosure_history = st.selectbox("Select foreclosure history:", ["None", "Within the past year", "1-5 years ago", "More than 5 years ago"], key="foreclosure_history")
    legal_issues = st.text_input("Enter any legal issues (pending lawsuits, liens, etc.):", key="legal_issues", value="None")

    temperature = 0.10
    max_output_tokens = 8192

    # Submit button
    generate_t2t = st.form_submit_button("Generate Report")

# Process the form submission
if generate_t2t:
    prompt = f"""
    You are an expert in underwriting and risk assessment for home loans. Using the input data provided, generate a comprehensive underwriting risk assessment report. The report should analyze the borrower's financial stability, creditworthiness, and the risk associated with the loan. Use the following structure for the report and make sure report generated only based on data provided and rules provided. Don't hallucinate.

    Applicant Information:
        Name: {applicant_name}
        Date of Birth: {date_of_birth}
        Social Security Number: {ssn}
        Marital Status: {marital_status}
        Number of Dependents: {number_of_dependents}

    Employment and Income:
        Current Employer: {current_employer}
        Job Title: {job_title}
        Years at Current Job: {years_at_current_job}
        Monthly Gross Income: {monthly_gross_income}
        Other Income Sources: {other_income_sources}

    Financial Information:
        Current Assets: {current_assets}
        Current Debts: {current_debts}
        Monthly Debt Obligations: {monthly_debt_obligations}
        Credit Score: {credit_score}

    Property Information:
        Property Address: {property_address}
        Property Type: {property_type}
        Purchase Price: {purchase_price}
        Loan Amount Requested: {loan_amount_requested}
        Down Payment Amount: {down_payment_amount}

    Loan Details:
        Loan Type: {loan_type}
        Loan Term: {loan_term}
        Interest Rate: {interest_rate}

    Housing Expense Information:
        Current Monthly Housing Payment: {current_monthly_housing_payment}
        Estimated Property Taxes: {estimated_property_taxes}
        Homeowners Insurance: {homeowners_insurance}
        HOA Fees: {hoa_fees}

    Additional Information:
        Purpose of Loan: {purpose_of_loan}
        Bankruptcy History: {bankruptcy_history}
        Foreclosure History: {foreclosure_history}
        Legal Issues: {legal_issues}

    Rules for Assessment:
    1. Loan-to-Value (LTV) Ratio Calculation:
       - LTV Ratio = (Loan Amount Requested / Purchase Price) * 100
       - High LTV (>80%) indicates higher risk.
       - Red Flag: LTV above 90%.

    2. Debt-to-Income (DTI) Ratio Calculation:
       - DTI Ratio = (Monthly Debt Obligations / Monthly Gross Income) * 100
       - A DTI above 43% is generally considered risky.
       - Red Flag: DTI above 50%.

    3. Credit Score Evaluation:
       - 750 and above: Excellent
       - 700-749: Good
       - 650-699: Fair
       - Below 650: Poor
       - Red Flag: Credit score below 600.

    4. Employment Stability:
       - Longer tenure at current job (>2 years) indicates stability.
       - Red Flag: Frequent job changes or employment less than 6 months.

    5. Current Assets vs. Debts:
       - Positive net worth (assets > debts) is favorable.
       - Red Flag: Net worth is negative.

    6. Housing Expense to Income Ratio:
       - Housing expenses should generally not exceed 28-31% of gross monthly income.
       - Red Flag: Housing expense ratio above 35%.

    Red Flags to Consider:
    - Recent bankruptcy (within the past 5 years).
    - History of foreclosure.
    - Legal issues such as pending lawsuits or liens.
    - Significant discrepancies between reported income and bank statements or tax returns.

    Sample Output:
    Underwriting Risk Assessment Report

    1. Applicant Information:
    - Name: John Doe
    - Date of Birth: January 1, 1980
    - Social Security Number: 123-45-6789
    - Marital Status: Married
    - Number of Dependents: 2

    2. Employment and Income:
    - Current Employer: XYZ Corporation
    - Job Title: Software Engineer
    - Years at Current Job: 5 years
    - Monthly Gross Income: $10,000
    - Other Income Sources: None

    3. Financial Summary:
    - Current Assets: $150,000
    - Current Debts: $50,000
    - Monthly Debt Obligations: $1,500
    - Credit Score: 720

    4. Property Details:
    - Property Address: 123 Elm Street, Springfield
    - Property Type: Single-family home
    - Purchase Price: $400,000
    - Loan Amount Requested: $320,000
    - Down Payment Amount: $80,000

    5. Loan Details:
    - Loan Type: Fixed-rate
    - Loan Term: 30 years
    - Interest Rate: 3.5%

    6. Housing Expense Analysis:
    - Current Monthly Housing Payment: $2,000 (rent)
    - Estimated Property Taxes: $4,000 annually
    - Homeowners Insurance: $1,200 annually
    - HOA Fees: None

    7. Risk Metrics:
    - LTV Ratio:
      - Calculation: ($320,000 / $400,000) * 100 = 80%
      - Assessment: Acceptable risk.

    - DTI Ratio:
      - Calculation: ($1,500 / $10,000) * 100 = 15%
      - Assessment: Low risk.

    - Credit Score:
      - Score: 720
      - Assessment: Good.

    - Employment Stability:
      - Years at Current Job: 5 years
      - Assessment: Stable.

    - Net Worth:
      - Calculation: $150,000 - $50,000 = $100,000
      - Assessment: Positive.

    - Housing Expense to Income Ratio:
      - Calculation: ($2,000 / $10,000) * 100 = 20%
      - Assessment: Affordable.

    8. Assessment and Recommendation:
    John Doe demonstrates a strong financial profile with a good credit score, stable employment, and manageable debt levels. The LTV and DTI ratios are within acceptable limits, indicating a moderate risk. However, key factors the underwriter should consider include:

    - Consistency of Income: Ensure that income is consistent with tax returns and bank statements.
    - Employment Verification: Confirm job stability and employment details with the employer.
    - Credit History: Review the detailed credit report for any late payments, collections, or other negative items not captured by the score.
    - Asset Verification: Verify the existence and liquidity of reported assets.

    Based on the provided information, it is recommended to approve the loan application for the requested amount of $320,000. No significant red flags were identified, and the applicant appears to be a low to moderate risk.
    """

    generation_config = GenerationConfig(
        temperature=temperature,
        max_output_tokens=max_output_tokens,
    )

    # Generate the report using the provided data
    with st.spinner("Generating report..."):
        generated_report = get_gemini_pro_text_response(
            model=text_model_pro,
            contents=prompt,
            generation_config=generation_config,
        )

    # remove special characters from output
    def clean_generated_report(report):
        cleaned_report = re.sub(r'[#$*]', '', report)   
        return cleaned_report
            
    cleaned_report = clean_generated_report(generated_report)
    
    # Function to convert Markdown to HTML with preserved formatting
    def markdown_to_html(text):
        html = text.replace('\n', '<br>')
        html = re.sub(r'(\*\*|__)(.*?)\1', r'<strong>\2</strong>', html)  # Bold text
        html = re.sub(r'(\*|_)(.*?)\1', r'<em>\2</em>', html)  # Italics text
        return html

# Convert the cleaned report to HTML
    html_report = markdown_to_html(cleaned_report)
        
    # Display the generated report
    # st.markdown("## Underwriting Risk Assessment Report")
    st.markdown(f"<div class='report-box'>{html_report}</div>", unsafe_allow_html=True)
    
   
    
    def save_report_as_word(report, file_path):
        doc = Document()
        doc.add_heading("Underwriting Risk Assessment Report", 0)
        doc.add_paragraph(report)
        doc.save(file_path)
        st.download_button(label="Download Report", data=open(file_path, 'rb'), file_name=file_path, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
     # Save and provide download link for the report as Word document
    report_file_path = f"Underwriting_Report_{applicant_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    
    
    
    save_report_as_word(cleaned_report, report_file_path)
    
    # st.markdown(f"[Download Report]({report_file_path})", unsafe_allow_html=True)
